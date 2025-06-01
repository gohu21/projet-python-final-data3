from docx import Document
from docx.shared import Inches

def generer_rapport(chemin="part2/output/rapport.docx"):
    doc = Document()

    doc.add_heading("Rapport - Projet Moby Dick", 0)

    doc.add_heading("1. Extraction du chapitre 1", level=1)
    doc.add_paragraph("Le chapitre 1 a été extrait automatiquement à partir du texte complet de Moby Dick.")

    doc.add_heading("2. Analyse des paragraphes", level=1)
    doc.add_paragraph("Chaque paragraphe du chapitre a été analysé pour en déterminer le nombre de mots. "
                      "Les longueurs ont été arrondies à la dizaine la plus proche.")

    doc.add_heading("3. Graphique de distribution", level=1)
    doc.add_paragraph("Voici la distribution des longueurs de paragraphes :")
    doc.add_picture("part2/output/graph.png", width=Inches(5))

    doc.add_heading("4. Création d'une image illustrée", level=1)
    doc.add_paragraph("Une image a été générée à partir d'une couverture avec ajout du titre, de l’auteur et du premier paragraphe.")

    doc.add_heading("5. Ajout d'un logo", level=1)
    doc.add_paragraph("Un logo noir et blanc a été inséré dans l’image précédente après rotation.")
    doc.add_picture("part2/output/image_avec_logo.jpg", width=Inches(5))

    doc.save(chemin)
    print(f"✅ Rapport généré : {chemin}")
